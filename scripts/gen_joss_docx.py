"""Preview of the JOSS paper as a docx.

JOSS itself builds the published PDF from paper.md via pandoc; this script
gives the authors a readable Word-format preview before submission. It
mirrors the markdown source (paper/joss_paper.md) section by section
with python-docx so the result is editable in Word for hand revisions.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

REPO = Path(
    "/mnt/c/Users/Giuseppe/OneDrive - Università Cattolica del Sacro Cuore/"
    "Metabolic Intelligence - Projects-MI/2024_QIMP/repo"
)
OUT = REPO / "paper" / "joss_paper_preview.docx"
FIG = REPO / "paper" / "figures"

SERIF = "Times New Roman"
MONO = "Consolas"
DARK = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)


def setup(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = SERIF
    n.font.size = Pt(11)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    n.paragraph_format.space_after = Pt(6)
    for level, size in [(1, 16), (2, 13), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = SERIF
        h.font.size = Pt(size)
        h.font.color.rgb = DARK
        h.font.bold = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)


def shade(cell, fill: str) -> None:
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(sh)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = MONO
    run.font.size = Pt(9.5)
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(sh)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, image_path: Path, caption: str, fig_n: int,
               width_in: float = 6.0) -> None:
    if not image_path.exists():
        add_para(doc, f"[missing image: {image_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cap.paragraph_format.space_after = Pt(8)
    cr = cap.add_run(f"Fig. {fig_n} ")
    cr.bold = True
    cr.font.size = Pt(10)
    tr = cap.add_run(caption)
    tr.font.size = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              caption: str, table_n: int) -> None:
    cap = doc.add_paragraph()
    cr = cap.add_run(f"Table {table_n}. ")
    cr.bold = True
    cr.font.size = Pt(10)
    tr = cap.add_run(caption)
    tr.font.size = Pt(10)
    tr.italic = True
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = SERIF
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(cell, "D9E1F2")
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.name = SERIF
    doc.add_paragraph()


# ---------------------------------------------------------- build --------


def build() -> Document:
    doc = Document()
    setup(doc)

    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("qimp-mi: a scalable Python library for Quantum Image Processing on Qiskit")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Software paper, prepared for submission to The Journal of Open Source Software (JOSS)")
    sr.italic = True
    sr.font.size = Pt(10)
    sr.font.color.rgb = GREY

    doc.add_paragraph()
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run("Author 1¹*, Author 2¹, Giulio Dolciami²").font.size = Pt(11)

    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    aff.add_run(
        "¹ Metabolic Intelligence Lab, Università Cattolica del Sacro Cuore, Milan, Italy\n"
        "² Department of Electronics and Telecommunications, Politecnico di Torino, Turin, Italy\n"
        "* Corresponding author: email@unicatt.it"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # ---- Summary -------------------------------------------------------
    doc.add_heading("Summary", level=1)
    add_para(doc,
        "qimp-mi is an open-source Python library that implements the canonical "
        "encodings, processing operations and benchmarking tools of the Quantum "
        "Image Processing (QIMP) field on top of Qiskit 2.x. The library covers "
        "the five image encodings that anchor the QIMP literature — FRQI (Le et al. "
        "2011), NEQR (Zhang et al. 2013), QPIE, MCRQI (Sun et al. 2013) and NCQI "
        "(Sang et al. 2017) — together with the standard processing operations "
        "catalogued in the Dolciami thesis (2022): geometric transforms (flip, "
        "swap, rotation, cyclic shift, region-restricted variants), chromatic "
        "operations (intensity complement, half-intensity, threshold classify, "
        "FRQI colour change), quantum arithmetic (ripple-carry adder, two's-"
        "complement subtractor, NEQR comparator), QHED edge detection, Quine–"
        "McCluskey circuit compression, a variational FRQI classifier, and a "
        "Generalized Polarization (GP) ratiometric circuit with a recently-"
        "derived closed-form parameter solution. All operations are parametrised "
        "over the spatial qubit count n and the intensity qubit count q; no fixed "
        "image size or bit depth is assumed.",
    )
    add_para(doc,
        "The library ships with a Streamlit-based interactive explorer that "
        "walks an image end-to-end through the load → preprocess → encode → "
        "process → execute pipeline, including direct submission to IBM Quantum "
        "hardware via the Sampler primitive. A unit suite of 338 tests covers "
        "every public function in a grid over n and q and verifies exact "
        "encoder/processor behaviour against statevector simulation. Type "
        "checking (mypy --strict), linting (ruff), API documentation "
        "(mkdocs-material with mkdocstrings), and packaging (PEP 621 with "
        "hatchling) are all in place; CI runs the test suite on three "
        "operating systems and three Python versions.",
    )

    # ---- Statement of need ---------------------------------------------
    doc.add_heading("Statement of need", level=1)
    add_para(doc,
        "Quantum image processing has accumulated more than twenty years of "
        "algorithm-level work — five competing encodings, dozens of processing "
        "operations, several application-specific extensions for multi-channel "
        "representation, compression and classification — surveyed in recent "
        "reviews (Yan & Venegas-Andraca 2025; Lisnichenko & Protasov 2023; "
        "Farooq et al. 2025). Despite this body of work, the field lacks a "
        "maintained, tested, scaling open-source library. Most published QIMP "
        "papers re-implement the FRQI or NEQR encoder from scratch, with "
        "hard-coded qubit counts and ad-hoc validation, often as supplementary "
        "code that is not kept up to date with the changing Qiskit API. This "
        "makes side-by-side comparison between papers difficult, reproducibility "
        "fragile, and the entry barrier for new researchers high.",
    )
    add_para(doc,
        "qimp-mi is intended to fill this gap. It provides the complete encoder "
        "/ processor catalog defined in the Dolciami thesis (2022), factored "
        "into composable, type-annotated primitives with parametric n / q "
        "everywhere, against the current (2.x) Qiskit API. Concretely, this "
        "serves three audiences:",
    )
    add_bullets(doc, [
        "New researchers can install the package and reproduce any of the "
        "standard QIMP operations on their own images in three lines of Python, "
        "rather than translating pseudocode from a paper.",
        "Methods researchers can side-by-side benchmark encodings on the same "
        "image — the bundled Benchmark page in the explorer runs FRQI, NEQR and "
        "QPIE on a single tile and emits a comparison table of qubit count, "
        "depth pre/post transpile, runtime and PSNR.",
        "Application researchers can take existing imaging operators "
        "(Generalized Polarization, NDVI, FRET, SNARF pH) and run them as "
        "quantum circuits with a closed-form parameter solver, without having "
        "to derive the parameters themselves.",
    ])
    add_para(doc,
        "A direct by-product of treating QIMP as a software-engineering target "
        "rather than a per-paper scratchpad was a new methodological result: a "
        "closed-form analytical solution for the parameters of a variational "
        "FRQI-based Generalized Polarization circuit, derived from a careful "
        "re-examination of the textbook ansatz. This result is reported in a "
        "companion methods paper (in preparation); the library ships the "
        "analytical_gp_params solver, the corrected apply_gp_function circuit, "
        "and a reproducible Streamlit validation page so that the derivation "
        "can be re-checked on any RGB tile.",
    )

    # ---- Library architecture ------------------------------------------
    doc.add_heading("Library architecture", level=1)
    add_figure(doc, FIG / "fig_qimp_architecture.png",
        "Architecture of the qimp-mi library. Top: the Streamlit application "
        "layer with its IBM Quantum Runtime hooks. Middle: the runtime / utility "
        "layer used by every other module. Library core: four problem-domain "
        "subpackages (encoding, processing, qml, io). Bottom: external "
        "dependencies. Each layer depends only on layers below it.",
        fig_n=1, width_in=6.0)
    add_para(doc,
        "The library is organised as a four-layer stack (Fig. 1). At the bottom "
        "sit the third-party dependencies (Qiskit 2.x, qiskit-aer, NumPy, "
        "SciPy, Pillow). Above them, the library core is split into four "
        "problem-domain subpackages (encoding, processing, qml, io). The third "
        "layer hosts the runtime, metrics and testing utilities shared across "
        "the core subpackages — simulator backend selection, performance "
        "monitoring, circuit caching, image metrics, noise-free and shot-based "
        "simulation, and IBM Quantum hardware execution. The top layer is the "
        "Streamlit application apps/qimp_explorer, which uses everything below "
        "it.",
    )

    # ---- Core abstractions ---------------------------------------------
    doc.add_heading("Core abstractions", level=1)
    add_para(doc,
        "The library exposes a small number of abstractions that compose "
        "uniformly across operations. The central one is the encoder / decoder "
        "pair: every encoding (FRQI, NEQR, QPIE, MCRQI, NCQI) is realised as a "
        "Python class with two methods,",
    )
    add_code(doc,
        "class XxxEncoder:\n"
        "    def encode(self, image: np.ndarray) -> QuantumCircuit: ...\n"
        "    def decode(self, counts: dict[str, int]) -> np.ndarray: ...",
    )
    add_para(doc,
        "so that a complete round-trip is decoded = enc.decode(simulate(enc."
        "encode(image))). Spatial and (where applicable) intensity qubit counts "
        "are inferred from the input image rather than configured ahead of time; "
        "image dtype and value range determine the FRQI normalisation, the NEQR "
        "bit width and the QPIE renormalisation automatically. The encoders "
        "deliberately return a circuit without a classical register attached, so "
        "that measurement is the caller's choice (deferred for downstream "
        "composition, or measure_all() for round-trip testing).",
    )
    add_para(doc,
        "Processing operations follow a complementary in-place convention:",
    )
    add_code(doc,
        "def axis_flip(qc: QuantumCircuit, n: int, *, axis: str,\n"
        "              pos_offset: int = 0) -> QuantumCircuit: ...",
    )
    add_para(doc,
        "Every operation takes the circuit, the spatial qubit count n, and an "
        "explicit pos_offset (the qubit index where the position register "
        "starts). Setting pos_offset to the encoding's intensity-register width "
        "makes the same geometric or chromatic operator work on a single- or "
        "multi-channel encoding without modification. The pattern keeps the "
        "processing layer agnostic to which encoder it is composed with.",
    )
    add_table(doc,
        headers=["Encoding", "Qubits", "Image type", "Exact?", "Channels"],
        rows=[
            ["FRQI", "2n + 1", "grayscale", "approx.*", "1"],
            ["NEQR", "2n + q", "grayscale", "exact",    "1"],
            ["QPIE", "2n",     "grayscale", "approx.*", "1"],
            ["MCRQI", "2n + 3", "RGB",      "approx.*", "3"],
            ["NCQI", "2n + 3q", "RGB",      "exact",    "3"],
        ],
        caption=(
            "The five image encodings provided by qimp.encoding. \"Exact\" means "
            "the decoded image matches the encoded image bit-for-bit given an "
            "exact simulator (NEQR and NCQI store intensities in computational-"
            "basis qubits; the others encode intensity as a rotation angle, so "
            "the decoded value is recovered up to shot noise)."
        ),
        table_n=1,
    )

    # ---- Subpackages ---------------------------------------------------
    doc.add_heading("Subpackages", level=1)

    doc.add_heading("qimp.encoding", level=2)
    add_para(doc,
        "Five encoder classes (FrqiEncoder, NeqrEncoder, QpieEncoder, "
        "McrqiEncoder, NcqiEncoder) plus a compression module that applies "
        "Quine–McCluskey logic minimisation with a greedy disjoint cover to "
        "reduce the number of multi-controlled gates needed for an FRQI / NEQR "
        "circuit. Compression is most effective on images with large constant "
        "regions (e.g. binary masks); the disjoint-cover step is needed because "
        "the standard QM output would let a single gate be controlled by "
        "overlapping minterms, which is correct for Boolean functions but wrong "
        "for the cumulative-rotation semantics of FRQI.",
    )

    doc.add_heading("qimp.processing", level=2)
    add_para(doc,
        "Five modules, each implementing a category of operation. geometric "
        "provides axis_flip, coord_swap, ort_rotation, pos_shift, restr_flip "
        "and restr_coord_swap. chromatic provides frqi_color_complement, "
        "frqi_color_change, neqr_color_complement, neqr_half_intensity and "
        "neqr_classify_complement. arithmetic provides qc_add_1, q_add, q_sub "
        "and neqr_comparator. filters provides QHED (the quantum Hadamard-edge-"
        "detection filter) and its two-pass full-edge variant. gp_ratio "
        "provides the corrected Generalized Polarization circuit "
        "(apply_gp_function), its decoder (decode_gp_counts), the classical "
        "reference image (classical_gp_image), the runnable end-to-end pipeline "
        "(evaluate_gp), the COBYLA wrapper (optimize_gp), and the closed-form "
        "parameter solver (analytical_gp_params).",
    )

    doc.add_heading("qimp.qml", level=2)
    add_para(doc,
        "The FrqiClassifier variational classifier wraps a FrqiEncoder with a "
        "RealAmplitudes ansatz (now via the functional "
        "qiskit.circuit.library.real_amplitudes API, with a class-form fallback "
        "for Qiskit 1.x users) and trains against ±1 labels via "
        "scipy.optimize.minimize(COBYLA). The classifier is intentionally "
        "minimal — a Hermitian-observable-on-the-color-qubit prediction with "
        "mean-squared-error loss — and serves as a working starting point "
        "rather than as a production-grade QML pipeline.",
    )

    doc.add_heading("qimp.io", level=2)
    add_para(doc,
        "TIFF loaders that preserve 16-bit depth, the classical GP image "
        "implementation calculate_gp_image with three output formats "
        "(normalized, uint8, 16bit), dataset iterators that walk an image "
        "directory and skip blank tiles, and image-buffer helpers "
        "(save_named_panels, new_output_dir) used by the Streamlit save "
        "buttons.",
    )

    doc.add_heading("Runtime / utility layer", level=2)
    add_para(doc,
        "qimp.runtime selects the Aer simulator backend (CPU or GPU) once per "
        "process, caches the FRQI scaffold circuit by (n, m), and provides a "
        "memory pool for image buffers in batch pipelines. qimp.metrics "
        "provides MSE, PSNR (with an explicit warning when a float image is "
        "supplied without max_intensity), total variation, and a "
        "transpile_summary helper used by the benchmarking page. qimp.testing "
        "provides three execution paths: ideal_simulation (noise-free shots), "
        "noisy_simulation (configurable noise model), and exact_counts "
        "(statevector → counts, deterministic, used in every regression test "
        "that asserts exact behaviour). It also wraps qiskit-ibm-runtime for "
        "hardware execution via the Sampler primitive.",
    )

    # ---- Application layer ---------------------------------------------
    doc.add_heading("Application layer", level=1)
    add_para(doc,
        "The Streamlit application apps/qimp_explorer is a five-step wizard "
        "(Load → Preprocess → Encode → Process → Execute / Export) plus three "
        "add-on pages: a Benchmark page that runs FRQI / NEQR / QPIE on the "
        "same image and emits a pandas comparison table, a GP-ratio page with "
        "the closed-form solver exposed as the primary action, and a System "
        "Info page that reports the simulator backend, dataset stats and "
        "past-run inventory. Circuit export to OpenQASM 3 and IBM Quantum "
        "hardware submission via the Sampler primitive are built in; the IBM "
        "token is held in st.session_state only and never persisted to disk.",
    )

    # ---- Example usage -------------------------------------------------
    doc.add_heading("Example usage", level=1)
    add_para(doc,
        "A complete classical-vs-quantum Generalized Polarization comparison "
        "on a microscopy tile is ten lines of Python:",
    )
    add_code(doc,
        "from PIL import Image\n"
        "import numpy as np\n"
        "from qimp.processing.gp_ratio import (\n"
        "    classical_gp_image, analytical_gp_params, evaluate_gp,\n"
        ")\n"
        "\n"
        "rgb = np.asarray(Image.open('membrane.tif').convert('RGB'))[:16, :16]\n"
        "I1, I2 = rgb[:, :, 0].astype(float), rgb[:, :, 1].astype(float)\n"
        "\n"
        "target  = classical_gp_image(I1, I2, alpha=0.5)\n"
        "params  = analytical_gp_params(I1, I2, alpha=0.5)\n"
        "decoded = evaluate_gp(I1, I2, params, exact=True)\n"
        "mse     = ((target - decoded) ** 2).mean()\n"
        "# → mse ≲ 1e-8 across realistic frames; PSNR ≥ 80 dB.",
    )
    add_para(doc,
        "The same circuit, with evaluate_gp(exact=False, shots=4096), runs on "
        "the shot-based simulator with no change to the parameter values. The "
        "Streamlit explorer reproduces this exact flow interactively.",
    )

    # ---- Validation and quality ----------------------------------------
    doc.add_heading("Validation and quality", level=1)
    add_para(doc,
        "qimp-mi is developed under a strict-quality regime: 338 unit tests "
        "parametrised over n ∈ {1..4} and q ∈ {1..3} (in line with the "
        "project's design constraint of scale-freeness), exact-statevector "
        "testing for the exact encoders (NEQR / NCQI / compression), regression "
        "tests for two architectural bugs in the original GP ansatz, and "
        "synthetic integration fixtures so the integration suite runs in CI "
        "without the lab's microscopy data being present. The CI pipeline runs "
        "on ubuntu-latest, macos-latest and windows-latest against Python 3.10, "
        "3.11 and 3.12; a separate build job runs python -m build && twine "
        "check dist/* to catch packaging regressions. Type checking is mypy "
        "--strict on the full src/qimp tree; linting is ruff check clean; "
        "documentation builds with mkdocs build --strict and includes "
        "auto-generated API reference via mkdocstrings.",
    )

    # ---- Comparison ----------------------------------------------------
    doc.add_heading("Comparison to existing software", level=1)
    add_para(doc,
        "To our knowledge no maintained library on PyPI implements the QIMP "
        "catalog at this scope. The closest neighbours are:",
    )
    add_bullets(doc, [
        "Qiskit itself, which provides the QuantumCircuit abstraction and the "
        "Aer simulator but no QIMP-specific encoders, processors or "
        "benchmarks. qimp-mi is built directly on Qiskit and Aer.",
        "PennyLane and TensorFlow Quantum target variational quantum machine "
        "learning broadly and have no first-class FRQI / NEQR support.",
        "qiskit-machine-learning provides VQC / QNN scaffolding but does not "
        "implement QIMP encodings; users wanting an FRQI-based classifier would "
        "have to write the encoder themselves.",
        "Numerous unmaintained research scripts on GitHub implement one or two "
        "encodings, usually with hard-coded qubit counts and no tests.",
    ])
    add_para(doc,
        "qimp-mi is complementary to all of these: it consumes Qiskit and Aer "
        "for the circuit primitives and simulator backend, and adds the "
        "QIMP-specific layer (encoders, processors, metrics, validation, "
        "interactive explorer, hardware hooks).",
    )

    # ---- Limitations and roadmap ---------------------------------------
    doc.add_heading("Limitations and roadmap", level=1)
    add_para(doc,
        "The current release (v0.3.0) covers the full Dolciami thesis catalogue "
        "and the GP closed-form result. Three items are explicitly on the "
        "roadmap: (i) per-channel normalisation for multi-channel ratiometric "
        "operators (the closed form already supports it via the normalization "
        "argument; the encoder needs to expose it); (ii) parameter-shift "
        "gradient-based training as an alternative to COBYLA for cost functions "
        "that are not per-pixel separable; (iii) a richer benchmark page that "
        "reports also the simulator wall-clock time per shot and the per-shot "
        "variance. The library targets simulation today; the IBM Quantum "
        "Runtime hook is exercised on small circuits but a systematic noise-"
        "budget study on hardware is left for a follow-up.",
    )

    # ---- Acknowledgements ----------------------------------------------
    doc.add_heading("Acknowledgements", level=1)
    add_para(doc,
        "We thank the members of the Metabolic Intelligence Lab for access to "
        "the fluorescence-microscopy dataset used to validate the GP pipeline, "
        "and Politecnico di Torino for the original thesis spec that defines "
        "the library's encoder / processor catalogue.",
    )

    doc.add_heading("References", level=1)
    add_para(doc,
        "References are managed in joss_paper.bib (BibTeX) and rendered by "
        "JOSS at build time from paper.md. See joss_paper.bib for the full "
        "list (11 entries covering FRQI, NEQR, MCRQI, NCQI, the three recent "
        "QIMP reviews, the Dolciami thesis, Qiskit, and the companion methods "
        "paper in preparation).",
    )

    return doc


def main() -> int:
    doc = build()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
